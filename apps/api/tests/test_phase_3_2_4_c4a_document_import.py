from __future__ import annotations

import base64
import struct
import zlib
from io import BytesIO
from pathlib import Path
from uuid import uuid4

import httpx
import pytest
from docx import Document
from docx.shared import Inches
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from sari_api.adapters.media_storage import LocalMediaStorage, get_media_storage
from sari_api.api.dependencies import get_token_identity
from sari_api.domain.identity import TokenIdentity
from sari_api.main import app

SALES_SUBJECT = "30000000-0000-4000-8000-000000000002"


async def sales_identity() -> TokenIdentity:
    return TokenIdentity(subject=SALES_SUBJECT, email="sales@sari-arta.example")


def png() -> bytes:
    def chunk(kind: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + kind
            + data
            + struct.pack(">I", zlib.crc32(kind + data))
        )

    ihdr = struct.pack(">IIBBBBB", 2, 3, 8, 2, 0, 0, 0)
    pixels = b"".join(b"\x00" + b"\x80\x80\x80" * 2 for _ in range(3))
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(
        b"IDAT", zlib.compress(pixels)
    ) + chunk(b"IEND", b"")


def docx_fixture(*, with_image: bool = False) -> bytes:
    document = Document()
    document.core_properties.title = "Synthetic School Kitchen Brief"
    document.add_heading("Project overview", level=1)
    document.add_paragraph("Plan preparation, cooking, washing and storage workflow.")
    document.add_paragraph("Confirm site conditions", style="List Bullet")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Facility"
    table.cell(0, 1).text = "School canteen"
    if with_image:
        document.add_picture(BytesIO(png()), width=Inches(1))
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def pdf_fixture(*, text: str | None) -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    if text:
        font = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        page[NameObject("/Resources")] = DictionaryObject(
            {
                NameObject("/Font"): DictionaryObject(
                    {NameObject("/F1"): writer._add_object(font)}
                )
            }
        )
        stream = DecodedStreamObject()
        safe = text.replace("(", "[").replace(")", "]")
        stream.set_data(f"BT /F1 12 Tf 72 720 Td ({safe}) Tj ET".encode())
        page[NameObject("/Contents")] = writer._add_object(stream)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


async def post_import(
    client: httpx.AsyncClient, filename: str, content: bytes, mime_type: str
) -> httpx.Response:
    return await client.post(
        "/api/v1/public-content/imports",
        files={"file": (filename, content, mime_type)},
    )


@pytest.mark.asyncio
async def test_docx_txt_and_pdf_import_preserve_structured_text(tmp_path: Path) -> None:
    app.dependency_overrides[get_token_identity] = sales_identity
    app.dependency_overrides[get_media_storage] = lambda: LocalMediaStorage(tmp_path / "media")
    transport = httpx.ASGITransport(app=app)
    try:
        async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
            docx = await post_import(
                client, "brief.docx", docx_fixture(with_image=True), DOCX_MIME
            )
            assert docx.status_code == 201, docx.text
            assert docx.json()["processing_status"] == "completed"
            assert docx.json()["extraction_result"]["title"] == "Synthetic School Kitchen Brief"
            kinds = {block["kind"] for block in docx.json()["extraction_result"]["blocks"]}
            assert {"heading", "paragraph", "list", "table"}.issubset(kinds)
            assert len(docx.json()["extracted_media_ids"]) == 1

            text = await post_import(
                client,
                "notes.txt",
                b"School canteen planning.\n\nConfirm washing and storage areas.",
                "text/plain",
            )
            assert text.status_code == 201
            assert text.json()["processing_status"] == "completed"
            assert len(text.json()["extraction_result"]["blocks"]) == 2

            pdf = await post_import(
                client,
                "brief.pdf",
                pdf_fixture(text="Synthetic central kitchen planning brief"),
                "application/pdf",
            )
            assert pdf.status_code == 201, pdf.text
            assert pdf.json()["processing_status"] == "completed"
            assert pdf.json()["extraction_result"]["blocks"][0]["page_number"] == 1
    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_html_is_structured_scripts_removed_and_remote_resources_ignored(
    tmp_path: Path,
) -> None:
    app.dependency_overrides[get_token_identity] = sales_identity
    app.dependency_overrides[get_media_storage] = lambda: LocalMediaStorage(tmp_path / "media")
    transport = httpx.ASGITransport(app=app)
    html = b"""
      <html><head><title>Safe synthetic brief</title><script>stealSecrets()</script></head>
      <body><h1>School canteen</h1><p>Review the kitchen workflow.</p>
      <img src="https://untrusted.example/image.png" alt="remote"></body></html>
    """
    try:
        async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
            response = await post_import(client, "brief.html", html, "text/html")
            assert response.status_code == 201
            body = response.json()
            assert body["processing_status"] == "completed"
            assert body["extraction_result"]["title"] == "Safe synthetic brief"
            extracted = " ".join(
                block["text"] for block in body["extraction_result"]["blocks"]
            )
            assert "School canteen" in extracted
            assert "stealSecrets" not in extracted
            assert body["extraction_metadata"]["removed_executable_elements"] == 1
            assert body["extraction_metadata"]["ignored_remote_images"] == 1
    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_embedded_image_becomes_private_media_and_import_creates_no_public_page(
    tmp_path: Path,
) -> None:
    app.dependency_overrides[get_token_identity] = sales_identity
    app.dependency_overrides[get_media_storage] = lambda: LocalMediaStorage(tmp_path / "media")
    transport = httpx.ASGITransport(app=app)
    data_uri = base64.b64encode(png()).decode()
    html = (
        "<h1>Synthetic brief</h1><p>Internal extraction fixture.</p>"
        f'<img alt="plan" src="data:image/png;base64,{data_uri}">'
    ).encode()
    try:
        async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
            before = await client.get("/api/v1/public-content/items")
            response = await post_import(client, "with-image.html", html, "text/html")
            assert response.status_code == 201, response.text
            body = response.json()
            assert body["processing_status"] == "completed"
            assert len(body["extracted_media_ids"]) == 1
            asset_id = body["extracted_media_ids"][0]
            asset = await client.get(f"/api/v1/media/assets/{asset_id}")
            assert asset.status_code == 200
            assert asset.json()["visibility"] == "private"
            assert asset.json()["public_use_status"] == "uploaded"
            assert asset.json()["source_reference_id"] == body["id"]
            assert (await client.get(f"/api/v1/media/public/{asset_id}")).status_code == 404
            after = await client.get("/api/v1/public-content/items")
            assert len(after.json()) == len(before.json())
    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_scanned_pdf_fails_safely_and_tenant_boundary_is_enforced(tmp_path: Path) -> None:
    app.dependency_overrides[get_token_identity] = sales_identity
    app.dependency_overrides[get_media_storage] = lambda: LocalMediaStorage(tmp_path / "media")
    transport = httpx.ASGITransport(app=app)
    try:
        async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
            response = await post_import(
                client, "scanned.pdf", pdf_fixture(text=None), "application/pdf"
            )
            assert response.status_code == 201
            body = response.json()
            assert body["processing_status"] == "failed"
            assert body["extraction_metadata"]["failure_code"] == "insufficient_extraction"
            assert body["extraction_result"] == {}
            isolated = await client.get(
                f"/api/v1/public-content/imports/{body['id']}",
                headers={"X-Tenant-Id": str(uuid4())},
            )
            assert isolated.status_code == 403
    finally:
        app.dependency_overrides.clear()


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
