from __future__ import annotations

import struct
import zlib
from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import UUID, uuid4

import httpx
import pytest
from docx import Document
from docx.shared import Inches
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from sari_api.adapters.media_storage import LocalMediaStorage, get_media_storage
from sari_api.adapters.public_content_structuring_provider import (
    MockPublicContentStructuringProvider,
)
from sari_api.api.dependencies import get_token_identity
from sari_api.api.routes.public_content_structuring import (
    get_public_content_structuring_provider,
)
from sari_api.domain.identity import TokenIdentity
from sari_api.domain.public_content_structuring import PublicContentStructuringResult
from sari_api.main import app

ADMIN_SUBJECT = "30000000-0000-4000-8000-000000000001"
SALES_SUBJECT = "30000000-0000-4000-8000-000000000002"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


async def admin_identity() -> TokenIdentity:
    return TokenIdentity(subject=ADMIN_SUBJECT, email="admin@sari-arta.example")


async def sales_identity() -> TokenIdentity:
    return TokenIdentity(subject=SALES_SUBJECT, email="sales@sari-arta.example")


async def unknown_identity() -> TokenIdentity:
    return TokenIdentity(subject=str(uuid4()), email="unknown@example.invalid")


class CountingProvider:
    provider_type = "mock"
    model_id = "counting-product-structuring-v1"

    def __init__(self) -> None:
        self.calls = 0
        self._delegate = MockPublicContentStructuringProvider()

    async def structure(
        self,
        *,
        import_id: UUID,
        import_result: dict[str, object],
        selected_page_type: str,
        locale: str,
    ) -> PublicContentStructuringResult:
        self.calls += 1
        return await self._delegate.structure(
            import_id=import_id,
            import_result=import_result,
            selected_page_type=selected_page_type,
            locale=locale,
        )


def png() -> bytes:
    def chunk(kind: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data)) + kind + data + struct.pack(">I", zlib.crc32(kind + data))
        )

    header = struct.pack(">IIBBBBB", 2, 3, 8, 2, 0, 0, 0)
    pixels = b"".join(b"\x00" + b"\x80\x80\x80" * 2 for _ in range(3))
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", header)
        + chunk(b"IDAT", zlib.compress(pixels))
        + chunk(b"IEND", b"")
    )


def product_docx() -> bytes:
    document = Document()
    document.core_properties.title = "Synthetic Product Source"
    document.add_heading("Synthetic Preparation Table", level=1)
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Product Name", "Synthetic Preparation Table"),
        ("SKU", "TEST-PT-180"),
        ("Category", "Preparation Equipment"),
        ("Material", "Stainless steel"),
        ("Dimensions", "1800 x 700 x 850 mm"),
        ("Power", "1 kW"),
        ("Price", "USD 180.00"),
        ("MOQ", "2 units"),
        ("Availability", "Made to order"),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.add_paragraph(
        "Short description: Synthetic source-backed preparation equipment description."
    )
    document.add_heading("Features", level=2)
    document.add_paragraph("Raised rear edge", style="List Bullet")
    document.add_heading("Applications", level=2)
    document.add_paragraph("Commercial kitchen preparation areas", style="List Bullet")
    document.add_picture(BytesIO(png()), width=Inches(1))
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def product_pdf() -> bytes:
    text = (
        "Product Name: Synthetic Sink Unit; SKU: TEST-SINK-01; "
        "Category: Washing Equipment; Material: Stainless steel; "
        "Dimensions: 1200 x 650 x 850 mm; "
        "Short description: Synthetic source-backed sink description"
    )
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)})}
    )
    stream = DecodedStreamObject()
    stream.set_data(f"BT /F1 10 Tf 36 720 Td ({text}) Tj ET".encode())
    page[NameObject("/Contents")] = writer._add_object(stream)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


async def upload(
    client: httpx.AsyncClient, filename: str, content: bytes, mime_type: str
) -> dict[str, Any]:
    response = await client.post(
        "/api/v1/public-content/imports",
        files={"file": (filename, content, mime_type)},
    )
    assert response.status_code == 201, response.text
    assert response.json()["processing_status"] == "completed"
    return response.json()


async def structure_product(client: httpx.AsyncClient, import_id: str) -> httpx.Response:
    return await client.post(
        f"/api/v1/public-content/imports/{import_id}/structure",
        json={"page_type": "product", "locale": "en"},
    )


@pytest.mark.asyncio
async def test_docx_product_structuring_preserves_facts_media_and_creates_draft_only(
    tmp_path: Path,
) -> None:
    app.dependency_overrides[get_token_identity] = sales_identity
    app.dependency_overrides[get_media_storage] = lambda: LocalMediaStorage(tmp_path / "media")
    transport = httpx.ASGITransport(app=app)
    try:
        async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
            imported = await upload(client, "product.docx", product_docx(), DOCX_MIME)
            response = await structure_product(client, imported["id"])
            assert response.status_code == 201, response.text
            run = response.json()
            assert run["selected_page_type"] == "product"
            assert run["provider"] == "mock"
            assert run["result"]["multiple_products_detected"] is False
            candidate = run["result"]["product_candidates"][0]
            content = candidate["content"]
            assert content["product_name"] == "Synthetic Preparation Table"
            assert content["sku_model"] == "TEST-PT-180"
            assert content["material"] == "Stainless steel"
            assert content["dimensions"] == "1800 x 700 x 850 mm"
            assert content["price_mode"] == "fixed"
            assert content["currency"] == "USD"
            assert content["price_min"] == "180.00"
            assert content["specifications"] == [{"label": "Power", "value": "1 kW"}]
            assert content["availability_note"] == "Made to order"
            media_id = imported["extracted_media_ids"][0]
            assert candidate["media_suggestions"][0]["media_asset_id"] == media_id
            assert content["hero_media_asset_id"] == media_id
            assert candidate["evidence"]

            draft = await client.post(
                f"/api/v1/public-content/imports/{imported['id']}/drafts",
                headers={"Idempotency-Key": f"product-draft-{uuid4()}"},
                json={
                    "structuring_run_id": run["id"],
                    "product_candidate_key": candidate["candidate_key"],
                    "slug": f"imported-product-{uuid4().hex[:8]}",
                    "title": candidate["title"],
                    "summary": candidate["summary"],
                    "seo_title": candidate["seo_title"],
                    "seo_description": candidate["seo_description"],
                    "structured_content": candidate["cms_structured_content"],
                    "media_references": [
                        {
                            "media_asset_id": media_id,
                            "role": "product_hero",
                            "alt_text": "Synthetic imported Product media pending review.",
                        }
                    ],
                },
            )
            assert draft.status_code == 201, draft.text
            item = draft.json()
            assert item["status"] == "draft"
            assert item["approved_version_id"] is None
            assert item["published_version_id"] is None
            assert item["current_version"]["origin"] == "ai_draft"
            assert item["current_version"]["source_structuring_run_id"] == run["id"]
            assert item["current_version"]["source_candidate_key"] == "product-1"
            assert item["current_version"]["source_reference_id"] == imported["id"]
            assert (await client.get(f"/api/v1/media/public/{media_id}")).status_code == 404

            app.dependency_overrides[get_token_identity] = admin_identity
            publish = await client.post(
                f"/api/v1/public-content/items/{item['id']}/publish",
                headers={"If-Match": '"1"', "Idempotency-Key": f"publish-{uuid4()}"},
                json={
                    "public_content_version_id": item["current_version"]["id"],
                    "content_sha256": item["current_version"]["content_sha256"],
                    "comment": "AI draft cannot bypass review or media approval.",
                },
            )
            assert publish.status_code == 409
    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_pdf_html_and_txt_are_deterministic_and_do_not_invent_price_or_currency(
    tmp_path: Path,
) -> None:
    app.dependency_overrides[get_token_identity] = sales_identity
    app.dependency_overrides[get_media_storage] = lambda: LocalMediaStorage(tmp_path / "media")
    transport = httpx.ASGITransport(app=app)
    try:
        async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
            pdf = await upload(client, "product.pdf", product_pdf(), "application/pdf")
            first = await structure_product(client, pdf["id"])
            second = await structure_product(client, pdf["id"])
            assert first.status_code == second.status_code == 201
            assert first.json()["result"] == second.json()["result"]
            content = first.json()["result"]["product_candidates"][0]["content"]
            assert content["sku_model"] == "TEST-SINK-01"
            assert content["material"] == "Stainless steel"
            assert content["dimensions"] == "1200 x 650 x 850 mm"
            assert content["price_mode"] == "request_quote"
            assert content["price_min"] is None
            assert content["price_max"] is None
            assert content["currency"] is None

            html = await upload(
                client,
                "product.html",
                b"<h1>Product</h1><p>Product Name: Synthetic Oven</p>"
                b"<p>SKU: TEST-OVEN</p><p>Category: Cooking Equipment</p>"
                b"<p>Short description: Synthetic HTML product source.</p>",
                "text/html",
            )
            text = await upload(
                client,
                "product.txt",
                b"Product Name: Synthetic Rack; SKU: TEST-RACK; "
                b"Category: Storage Equipment; Short description: Synthetic TXT source.",
                "text/plain",
            )
            markdown = await upload(
                client,
                "product.md",
                b"# Synthetic Shelf\n\nProduct Name: Synthetic Shelf; SKU: TEST-SHELF; "
                b"Category: Storage Equipment; Short description: Synthetic Markdown source.",
                "text/markdown",
            )
            assert (await structure_product(client, html["id"])).status_code == 201
            assert (await structure_product(client, text["id"])).status_code == 201
            assert (await structure_product(client, markdown["id"])).status_code == 201
    finally:
        app.dependency_overrides.clear()


@pytest.mark.asyncio
async def test_multiple_products_require_selection_and_authorization_precedes_provider(
    tmp_path: Path,
) -> None:
    provider = CountingProvider()
    app.dependency_overrides[get_token_identity] = sales_identity
    app.dependency_overrides[get_media_storage] = lambda: LocalMediaStorage(tmp_path / "media")
    app.dependency_overrides[get_public_content_structuring_provider] = lambda: provider
    transport = httpx.ASGITransport(app=app)
    html = b"""
      <h1>Product A</h1><p>Product Name: Synthetic Table A</p><p>SKU: TEST-A</p>
      <p>Category: Preparation Equipment</p><p>Short description: First product source.</p>
      <h1>Product B</h1><p>Product Name: Synthetic Table B</p><p>SKU: TEST-B</p>
      <p>Category: Preparation Equipment</p><p>Short description: Second product source.</p>
    """
    try:
        async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
            imported = await upload(client, "catalog.html", html, "text/html")

            app.dependency_overrides[get_token_identity] = unknown_identity
            denied = await structure_product(client, imported["id"])
            assert denied.status_code == 403
            assert provider.calls == 0

            app.dependency_overrides[get_token_identity] = sales_identity
            isolated = await client.post(
                f"/api/v1/public-content/imports/{imported['id']}/structure",
                headers={"X-Tenant-Id": str(uuid4())},
                json={"page_type": "product", "locale": "en"},
            )
            assert isolated.status_code == 403
            assert provider.calls == 0

            response = await structure_product(client, imported["id"])
            assert response.status_code == 201, response.text
            assert provider.calls == 1
            run = response.json()
            assert run["result"]["multiple_products_detected"] is True
            candidates = run["result"]["product_candidates"]
            assert [candidate["content"]["sku_model"] for candidate in candidates] == [
                "TEST-A",
                "TEST-B",
            ]

            first = candidates[0]
            blocked = await client.post(
                f"/api/v1/public-content/imports/{imported['id']}/drafts",
                headers={"Idempotency-Key": f"missing-candidate-{uuid4()}"},
                json={
                    "structuring_run_id": run["id"],
                    "slug": f"multiple-product-{uuid4().hex[:8]}",
                    "title": first["title"],
                    "summary": first["summary"],
                    "seo_title": first["seo_title"],
                    "seo_description": first["seo_description"],
                    "structured_content": first["cms_structured_content"],
                    "media_references": [],
                },
            )
            assert blocked.status_code == 409
    finally:
        app.dependency_overrides.clear()
