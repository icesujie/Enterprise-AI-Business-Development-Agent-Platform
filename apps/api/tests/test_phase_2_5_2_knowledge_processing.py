from __future__ import annotations

from io import BytesIO
from pathlib import Path
from uuid import UUID, uuid4

import httpx
import pytest
from docx import Document
from sqlalchemy import select, text

from sari_api.adapters.database import session_factory
from sari_api.adapters.knowledge_embedding import DeterministicKnowledgeEmbeddingProvider
from sari_api.adapters.knowledge_extractor import DefaultKnowledgeTextExtractor
from sari_api.adapters.knowledge_processing_queue import get_knowledge_processing_queue
from sari_api.adapters.knowledge_storage import LocalKnowledgeStorage, get_knowledge_storage
from sari_api.adapters.managed_knowledge_processing import ManagedKnowledgeProcessingExecutor
from sari_api.adapters.models import ManagedKnowledgeChunk
from sari_api.api.dependencies import get_token_identity
from sari_api.domain.identity import TokenIdentity
from sari_api.main import app

ADMIN_SUBJECT = "30000000-0000-4000-8000-000000000001"
TENANT_ID = UUID("10000000-0000-4000-8000-000000000001")


async def admin_identity() -> TokenIdentity:
    return TokenIdentity(subject=ADMIN_SUBJECT, email="admin@sari-arta.example")


class RecordingProcessingQueue:
    def __init__(self) -> None:
        self.messages: list[tuple[UUID, UUID, str | None]] = []

    async def enqueue(
        self,
        processing_run_id: UUID,
        tenant_id: UUID,
        *,
        correlation_id: str | None = None,
    ) -> None:
        self.messages.append((processing_run_id, tenant_id, correlation_id))


def docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Facility requirements", level=1)
    document.add_paragraph("Synthetic IVC facility capacity is 120 cages.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Timeline"
    table.rows[0].cells[1].text = "Q2 2027"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.mark.asyncio
async def test_docx_extraction_preserves_heading_and_table() -> None:
    sections = await DefaultKnowledgeTextExtractor().extract(
        docx_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert sections[0].section_title == "Facility requirements"
    assert "120 cages" in sections[0].text
    assert "Timeline | Q2 2027" in sections[0].text


@pytest.mark.asyncio
async def test_approved_document_processing_creates_isolated_citable_chunks(
    tmp_path: Path,
) -> None:
    storage = LocalKnowledgeStorage(tmp_path / "knowledge")
    queue = RecordingProcessingQueue()
    app.dependency_overrides[get_token_identity] = admin_identity
    app.dependency_overrides[get_knowledge_storage] = lambda: storage
    app.dependency_overrides[get_knowledge_processing_queue] = lambda: queue
    suffix = uuid4().hex[:10]
    collection_id: str | None = None
    try:
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(transport=transport, base_url="http://test") as client:
            collection = await client.post(
                "/api/v1/knowledge-management/collections",
                json={
                    "domain_key": "laboratory_animal_facility",
                    "collection_key": f"processing-{suffix}",
                    "name": f"Synthetic Processing {suffix}",
                },
            )
            assert collection.status_code == 201, collection.text
            collection_id = collection.json()["id"]
            upload = await client.post(
                f"/api/v1/knowledge-management/collections/{collection_id}/documents",
                data={
                    "title": f"Synthetic DOCX {suffix}",
                    "document_type": "technical_reference",
                    "language": "en",
                    "document_metadata_json": '{"synthetic":true,"region":"demo"}',
                },
                files={
                    "file": (
                        "synthetic-ivc.docx",
                        docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
            assert upload.status_code == 201, upload.text
            document_id = upload.json()["id"]
            assert upload.json()["processing_status"] == "uploaded"
            binding = await client.post(
                f"/api/v1/knowledge-management/documents/{document_id}/bindings",
                json={"agent_key": "laboratory_animal_facility.ivc_business_development"},
            )
            assert binding.status_code == 201, binding.text

            denied = await client.post(
                f"/api/v1/knowledge-management/documents/{document_id}/processing-runs"
            )
            assert denied.status_code == 409

            submitted = await client.post(
                f"/api/v1/knowledge-management/documents/{document_id}/submit-review"
            )
            assert submitted.status_code == 200
            approved = await client.post(
                f"/api/v1/knowledge-management/documents/{document_id}/approval",
                json={"decision": "approved"},
            )
            assert approved.status_code == 200

            processing = await client.post(
                f"/api/v1/knowledge-management/documents/{document_id}/processing-runs",
                headers={"X-Correlation-ID": "processing-test-001"},
            )
            assert processing.status_code == 202, processing.text
            run_id = UUID(processing.json()["id"])
            assert processing.json()["status"] == "uploaded"
            assert queue.messages == [(run_id, TENANT_ID, "processing-test-001")]

            await ManagedKnowledgeProcessingExecutor(
                storage,
                DeterministicKnowledgeEmbeddingProvider(1536),
            ).execute(run_id, TENANT_ID)

            result = await client.get(f"/api/v1/knowledge-management/processing-runs/{run_id}")
            assert result.status_code == 200, result.text
            assert result.json()["status"] == "completed"
            assert result.json()["chunk_count"] >= 1
            detail = await client.get(f"/api/v1/knowledge-management/documents/{document_id}")
            assert detail.json()["processing_status"] == "completed"

            async with session_factory() as session:
                await session.execute(
                    text("SELECT set_config('app.tenant_id', :tenant_id, true)"),
                    {"tenant_id": str(TENANT_ID)},
                )
                chunks = list(
                    (
                        await session.scalars(
                            select(ManagedKnowledgeChunk).where(
                                ManagedKnowledgeChunk.document_id == UUID(document_id),
                                ManagedKnowledgeChunk.tenant_id == TENANT_ID,
                            )
                        )
                    ).all()
                )
            assert chunks
            chunk = chunks[0]
            assert chunk.document_version_id == UUID(processing.json()["document_version_id"])
            assert chunk.citation_metadata["document_id"] == document_id
            assert chunk.citation_metadata["filename"] == "synthetic-ivc.docx"
            assert chunk.citation_metadata["section_title"] == "Facility requirements"
            assert chunk.source_metadata["document_metadata"]["region"] == "demo"
            assert len(chunk.embedding) == 1536

            stored_file = next((tmp_path / "knowledge").rglob("*.docx"))
            stored_file.unlink()
            retry = await client.post(
                f"/api/v1/knowledge-management/documents/{document_id}/processing-runs"
            )
            assert retry.status_code == 202, retry.text
            failed_run_id = UUID(retry.json()["id"])
            await ManagedKnowledgeProcessingExecutor(
                storage,
                DeterministicKnowledgeEmbeddingProvider(1536),
            ).execute(failed_run_id, TENANT_ID)
            failed = await client.get(
                f"/api/v1/knowledge-management/processing-runs/{failed_run_id}"
            )
            assert failed.json()["status"] == "failed"
            assert failed.json()["error_code"] == "object_missing"
            failed_detail = await client.get(
                f"/api/v1/knowledge-management/documents/{document_id}"
            )
            assert failed_detail.json()["processing_status"] == "failed"
    finally:
        app.dependency_overrides.clear()
        if collection_id:
            async with session_factory() as cleanup:
                await cleanup.execute(
                    text("SELECT set_config('app.tenant_id', :tenant_id, true)"),
                    {"tenant_id": str(TENANT_ID)},
                )
                for statement in (
                    "DELETE FROM managed_knowledge_chunks WHERE collection_id = :id",
                    "DELETE FROM knowledge_processing_runs WHERE document_id IN "
                    "(SELECT id FROM managed_knowledge_documents WHERE collection_id = :id)",
                    "DELETE FROM knowledge_audit_logs WHERE document_id IN "
                    "(SELECT id FROM managed_knowledge_documents WHERE collection_id = :id)",
                    "DELETE FROM knowledge_document_agent_bindings WHERE document_id IN "
                    "(SELECT id FROM managed_knowledge_documents WHERE collection_id = :id)",
                    "UPDATE managed_knowledge_documents SET current_version_id = NULL, "
                    "published_version_id = NULL, active_version_id = NULL "
                    "WHERE collection_id = :id",
                    "DELETE FROM knowledge_document_versions WHERE document_id IN "
                    "(SELECT id FROM managed_knowledge_documents WHERE collection_id = :id)",
                    "DELETE FROM managed_knowledge_documents WHERE collection_id = :id",
                    "DELETE FROM knowledge_collections WHERE id = :id",
                ):
                    await cleanup.execute(text(statement), {"id": collection_id})
                await cleanup.commit()


@pytest.mark.asyncio
async def test_processing_tables_force_rls_and_cross_tenant_reads_are_empty() -> None:
    async with session_factory() as session:
        rows = (
            await session.execute(
                text(
                    "SELECT relname, relrowsecurity, relforcerowsecurity FROM pg_class "
                    "WHERE relname IN ('knowledge_processing_runs','managed_knowledge_chunks') "
                    "ORDER BY relname"
                )
            )
        ).all()
        other_tenant = uuid4()
        await session.execute(
            text("SELECT set_config('app.tenant_id', :tenant_id, true)"),
            {"tenant_id": str(other_tenant)},
        )
        visible = await session.scalar(
            select(ManagedKnowledgeChunk.id).where(ManagedKnowledgeChunk.tenant_id == other_tenant)
        )
    assert len(rows) == 2
    assert all(row.relrowsecurity and row.relforcerowsecurity for row in rows)
    assert visible is None
