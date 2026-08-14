from __future__ import annotations

import asyncio
from io import BytesIO
from typing import Protocol

from docx import Document
from pypdf import PdfReader

from sari_api.domain.knowledge import ExtractedSection, normalize_extracted_text


class UnsupportedKnowledgeDocumentError(Exception):
    pass


class EmptyKnowledgeDocumentError(Exception):
    pass


class KnowledgeTextExtractor(Protocol):
    async def extract(self, content: bytes, media_type: str) -> list[ExtractedSection]: ...


class DefaultKnowledgeTextExtractor:
    supported_media_types = frozenset(
        {
            "application/pdf",
            "text/plain",
            "text/markdown",
            "text/x-markdown",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
    )

    async def extract(self, content: bytes, media_type: str) -> list[ExtractedSection]:
        normalized_type = media_type.split(";", maxsplit=1)[0].strip().lower()
        if normalized_type not in self.supported_media_types:
            raise UnsupportedKnowledgeDocumentError
        if normalized_type == "application/pdf":
            sections = await asyncio.to_thread(self._extract_pdf, content)
        elif normalized_type == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            sections = await asyncio.to_thread(self._extract_docx, content)
        elif normalized_type in {"text/markdown", "text/x-markdown"}:
            sections = self._extract_markdown(content)
        else:
            sections = self._extract_text(content)
        if not any(section.text.strip() for section in sections):
            raise EmptyKnowledgeDocumentError
        return sections

    @staticmethod
    def _extract_text(content: bytes) -> list[ExtractedSection]:
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise UnsupportedKnowledgeDocumentError from exc
        return [ExtractedSection(text=normalize_extracted_text(text))]

    @staticmethod
    def _extract_pdf(content: bytes) -> list[ExtractedSection]:
        try:
            reader = PdfReader(BytesIO(content))
            return [
                ExtractedSection(
                    text=normalize_extracted_text(page.extract_text() or ""),
                    page_number=index,
                )
                for index, page in enumerate(reader.pages, start=1)
            ]
        except Exception as exc:
            raise UnsupportedKnowledgeDocumentError from exc

    @staticmethod
    def _extract_markdown(content: bytes) -> list[ExtractedSection]:
        try:
            raw = content.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise UnsupportedKnowledgeDocumentError from exc
        sections: list[ExtractedSection] = []
        heading: str | None = None
        body: list[str] = []
        for line in raw.splitlines():
            if line.lstrip().startswith("#"):
                if body:
                    sections.append(
                        ExtractedSection(
                            text=normalize_extracted_text("\n".join(body)),
                            section_title=heading,
                        )
                    )
                heading = line.lstrip("#").strip() or None
                body = []
            else:
                body.append(line)
        if body or heading:
            sections.append(
                ExtractedSection(
                    text=normalize_extracted_text("\n".join(body)),
                    section_title=heading,
                )
            )
        return sections

    @staticmethod
    def _extract_docx(content: bytes) -> list[ExtractedSection]:
        try:
            document = Document(BytesIO(content))
            sections: list[ExtractedSection] = []
            heading: str | None = None
            body: list[str] = []
            for paragraph in document.paragraphs:
                value = normalize_extracted_text(paragraph.text)
                if not value:
                    continue
                style_name = paragraph.style.name if paragraph.style else ""
                if style_name.casefold().startswith("heading"):
                    if body:
                        sections.append(
                            ExtractedSection(
                                text=normalize_extracted_text("\n".join(body)),
                                section_title=heading,
                            )
                        )
                    heading = value
                    body = []
                else:
                    body.append(value)
            for table in document.tables:
                rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
                body.extend(value for value in rows if value.strip(" |"))
            if body or heading:
                sections.append(
                    ExtractedSection(
                        text=normalize_extracted_text("\n".join(body)),
                        section_title=heading,
                    )
                )
            return sections
        except Exception as exc:
            raise UnsupportedKnowledgeDocumentError from exc
