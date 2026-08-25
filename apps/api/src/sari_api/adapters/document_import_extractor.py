from __future__ import annotations

import base64
import binascii
import re
from collections.abc import Sequence
from dataclasses import dataclass
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from typing import ClassVar, Literal

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader

from sari_api.adapters.knowledge_extractor import (
    DefaultKnowledgeTextExtractor,
    EmptyKnowledgeDocumentError,
    UnsupportedKnowledgeDocumentError,
)
from sari_api.domain.knowledge import ExtractedSection, normalize_extracted_text


class UnsupportedImportDocumentError(ValueError):
    pass


class InsufficientImportExtractionError(ValueError):
    pass


@dataclass(frozen=True, slots=True)
class ImportBlock:
    kind: Literal["heading", "paragraph", "list", "table"]
    text: str
    order: int
    level: int | None = None
    page_number: int | None = None
    section_title: str | None = None


@dataclass(frozen=True, slots=True)
class ImportImage:
    content: bytes
    filename: str
    mime_type: str
    order: int
    page_number: int | None = None
    section_title: str | None = None


@dataclass(frozen=True, slots=True)
class DocumentImportResult:
    title: str | None
    blocks: list[ImportBlock]
    images: list[ImportImage]
    metadata: dict[str, object]


class DefaultDocumentImportExtractor:
    """Structured import adapter that reuses the Knowledge text extractor."""

    supported_media_types = frozenset(
        {
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "text/html",
            "text/plain",
            "text/markdown",
            "text/x-markdown",
        }
    )

    def __init__(self) -> None:
        self._knowledge = DefaultKnowledgeTextExtractor()

    async def extract(self, content: bytes, media_type: str) -> DocumentImportResult:
        normalized = media_type.split(";", maxsplit=1)[0].strip().lower()
        if normalized not in self.supported_media_types:
            raise UnsupportedImportDocumentError("Unsupported document type.")
        try:
            if normalized == "text/html":
                result = self._extract_html(content)
            elif normalized == "application/pdf":
                sections = await self._knowledge.extract(content, normalized)
                result = self._extract_pdf(content, sections)
            elif normalized == (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                await self._knowledge.extract(content, normalized)
                result = self._extract_docx(content)
            elif normalized in {"text/markdown", "text/x-markdown"}:
                sections = await self._knowledge.extract(content, normalized)
                result = DocumentImportResult(
                    title=next(
                        (section.section_title for section in sections if section.section_title),
                        None,
                    ),
                    blocks=[
                        ImportBlock(
                            kind="paragraph",
                            text=section.text,
                            order=index,
                            section_title=section.section_title,
                        )
                        for index, section in enumerate(sections)
                        if section.text
                    ],
                    images=[],
                    metadata={"format": "markdown", "extractor": "knowledge_extractor_v1"},
                )
            else:
                sections = await self._knowledge.extract(content, normalized)
                paragraphs = [
                    normalize_extracted_text(value)
                    for value in sections[0].text.split("\n\n")
                    if normalize_extracted_text(value)
                ]
                result = DocumentImportResult(
                    title=None,
                    blocks=[
                        ImportBlock(kind="paragraph", text=value, order=index)
                        for index, value in enumerate(paragraphs)
                    ],
                    images=[],
                    metadata={"format": "text", "extractor": "knowledge_extractor_v1"},
                )
        except EmptyKnowledgeDocumentError as exc:
            raise InsufficientImportExtractionError(
                "No usable text was extracted. Scanned PDFs require OCR, which is not supported."
            ) from exc
        except UnsupportedKnowledgeDocumentError as exc:
            raise UnsupportedImportDocumentError(
                "The document could not be safely parsed."
            ) from exc
        if not any(block.text.strip() for block in result.blocks):
            raise InsufficientImportExtractionError(
                "No usable text was extracted. Scanned PDFs require OCR, which is not supported."
            )
        return result

    @staticmethod
    def _extract_pdf(
        content: bytes, sections: Sequence[ExtractedSection]
    ) -> DocumentImportResult:
        try:
            reader = PdfReader(BytesIO(content))
            metadata = reader.metadata
            title = str(metadata.title).strip() if metadata and metadata.title else None
            blocks: list[ImportBlock] = []
            images: list[ImportImage] = []
            warnings: list[str] = []
            for index, section in enumerate(sections):
                text = normalize_extracted_text(str(getattr(section, "text", "")))
                if text:
                    blocks.append(
                        ImportBlock(
                            kind="paragraph",
                            text=text,
                            order=len(blocks),
                            page_number=index + 1,
                        )
                    )
            for page_number, page in enumerate(reader.pages, start=1):
                try:
                    for image in page.images:
                        filename = Path(image.name or f"page-{page_number}-image").name
                        mime_type = _mime_for_filename(filename)
                        if mime_type:
                            images.append(
                                ImportImage(
                                    content=image.data,
                                    filename=filename,
                                    mime_type=mime_type,
                                    order=len(images),
                                    page_number=page_number,
                                )
                            )
                except Exception:
                    warnings.append(f"Images on PDF page {page_number} could not be extracted.")
            return DocumentImportResult(
                title=title,
                blocks=blocks,
                images=images,
                metadata={
                    "format": "pdf",
                    "page_count": len(reader.pages),
                    "extractor": "knowledge_extractor_v1",
                    "warnings": warnings,
                },
            )
        except InsufficientImportExtractionError:
            raise
        except Exception as exc:
            raise UnsupportedImportDocumentError("The PDF could not be safely parsed.") from exc

    @staticmethod
    def _extract_docx(content: bytes) -> DocumentImportResult:
        try:
            document = Document(BytesIO(content))
            blocks: list[ImportBlock] = []
            current_section: str | None = None
            image_locations: dict[str, tuple[int, str | None]] = {}
            for item in document.iter_inner_content():
                if isinstance(item, Paragraph):
                    text = normalize_extracted_text(item.text)
                    style = item.style.name if item.style else ""
                    if text and style.casefold().startswith("heading"):
                        level_match = re.search(r"(\d+)", style)
                        level = int(level_match.group(1)) if level_match else 2
                        current_section = text
                        kind: Literal["heading", "paragraph", "list", "table"] = "heading"
                    elif text and "list" in style.casefold():
                        level = None
                        kind = "list"
                    else:
                        level = None
                        kind = "paragraph"
                    if text:
                        blocks.append(
                            ImportBlock(
                                kind=kind,
                                text=text,
                                order=len(blocks),
                                level=level,
                                section_title=current_section,
                            )
                        )
                    for relationship_id in item._p.xpath(".//a:blip/@r:embed"):
                        image_locations[str(relationship_id)] = (
                            len(blocks),
                            current_section,
                        )
                elif isinstance(item, Table):
                    rows = [
                        " | ".join(cell.text.strip() for cell in row.cells)
                        for row in item.rows
                    ]
                    value = normalize_extracted_text(
                        "\n".join(row for row in rows if row.strip(" |"))
                    )
                    if not value:
                        continue
                    blocks.append(
                        ImportBlock(
                            kind="table",
                            text=value,
                            order=len(blocks),
                            section_title=current_section,
                        )
                    )
            images: list[ImportImage] = []
            ordered_relationships = sorted(
                document.part.rels.items(),
                key=lambda entry: image_locations.get(entry[0], (len(blocks) + 1, None))[0],
            )
            for relationship_id, relationship in ordered_relationships:
                if relationship.reltype != RELATIONSHIP_TYPE.IMAGE:
                    continue
                part = relationship.target_part
                filename = Path(str(part.partname)).name
                mime_type = str(part.content_type)
                if mime_type not in {"image/jpeg", "image/png", "image/webp"}:
                    continue
                images.append(
                    ImportImage(
                        content=part.blob,
                        filename=filename,
                        mime_type=mime_type,
                        order=len(images),
                        section_title=image_locations.get(relationship_id, (0, None))[1],
                    )
                )
            title = document.core_properties.title.strip() or None
            return DocumentImportResult(
                title=title,
                blocks=blocks,
                images=images,
                metadata={
                    "format": "docx",
                    "extractor": "knowledge_extractor_v1",
                    "paragraph_count": len(document.paragraphs),
                    "table_count": len(document.tables),
                },
            )
        except Exception as exc:
            raise UnsupportedImportDocumentError(
                "The DOCX file could not be safely parsed."
            ) from exc

    @staticmethod
    def _extract_html(content: bytes) -> DocumentImportResult:
        try:
            raw = content.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise UnsupportedImportDocumentError("HTML imports must use UTF-8.") from exc
        parser = _SafeImportHTMLParser()
        try:
            parser.feed(raw)
            parser.close()
        except Exception as exc:
            raise UnsupportedImportDocumentError(
                "The HTML file could not be safely parsed."
            ) from exc
        return DocumentImportResult(
            title=parser.title,
            blocks=parser.blocks,
            images=parser.images,
            metadata={
                "format": "html",
                "extractor": "safe_html_parser_v1",
                "ignored_remote_images": parser.ignored_remote_images,
                "removed_executable_elements": parser.removed_executable_elements,
            },
        )


class _SafeImportHTMLParser(HTMLParser):
    _block_tags: ClassVar[dict[str, Literal["paragraph", "list", "table"]]] = {
        "p": "paragraph",
        "li": "list",
        "tr": "table",
    }

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[ImportBlock] = []
        self.images: list[ImportImage] = []
        self.title: str | None = None
        self.ignored_remote_images = 0
        self.removed_executable_elements = 0
        self._ignored_depth = 0
        self._capture_tag: str | None = None
        self._capture_kind: Literal["heading", "paragraph", "list", "table"] | None = None
        self._capture_level: int | None = None
        self._text: list[str] = []
        self._current_section: str | None = None
        self._in_title = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.casefold()
        if tag in {"script", "style", "iframe", "object", "embed", "svg"}:
            self.removed_executable_elements += 1
            self._ignored_depth += 1
            return
        if self._ignored_depth:
            return
        if tag == "title":
            self._in_title = True
            self._capture_tag = tag
            self._text = []
        elif re.fullmatch(r"h[1-6]", tag):
            self._begin(tag, "heading", int(tag[1]))
        elif tag in self._block_tags:
            self._begin(tag, self._block_tags[tag], None)
        elif tag == "img":
            values = dict(attrs)
            self._image(values.get("src"), values.get("alt"))

    def handle_endtag(self, tag: str) -> None:
        tag = tag.casefold()
        if tag in {"script", "style", "iframe", "object", "embed", "svg"}:
            if self._ignored_depth:
                self._ignored_depth -= 1
            return
        if self._ignored_depth:
            return
        if tag == "title" and self._in_title:
            self.title = normalize_extracted_text(" ".join(self._text)) or None
            self._in_title = False
            self._capture_tag = None
            self._text = []
        elif tag == self._capture_tag:
            self._finish()

    def handle_data(self, data: str) -> None:
        if not self._ignored_depth and self._capture_tag:
            self._text.append(data)

    def _begin(
        self,
        tag: str,
        kind: Literal["heading", "paragraph", "list", "table"],
        level: int | None,
    ) -> None:
        if self._capture_tag and not self._in_title:
            self._finish()
        self._capture_tag = tag
        self._capture_kind = kind
        self._capture_level = level
        self._text = []

    def _finish(self) -> None:
        value = normalize_extracted_text(" ".join(self._text))
        if value and self._capture_kind:
            if self._capture_kind == "heading":
                self._current_section = value
            self.blocks.append(
                ImportBlock(
                    kind=self._capture_kind,
                    text=value,
                    order=len(self.blocks),
                    level=self._capture_level,
                    section_title=self._current_section,
                )
            )
        self._capture_tag = None
        self._capture_kind = None
        self._capture_level = None
        self._text = []

    def _image(self, src: str | None, alt: str | None) -> None:
        if not src or not src.startswith("data:image/"):
            self.ignored_remote_images += 1
            return
        match = re.fullmatch(r"data:(image/(?:png|jpeg|webp));base64,(.+)", src, re.DOTALL)
        if not match:
            return
        try:
            content = base64.b64decode(match.group(2), validate=True)
        except (binascii.Error, ValueError):
            return
        extension = {"image/png": ".png", "image/jpeg": ".jpg", "image/webp": ".webp"}[
            match.group(1)
        ]
        safe_stem = re.sub(r"[^a-z0-9]+", "-", (alt or "embedded-image").casefold()).strip("-")
        self.images.append(
            ImportImage(
                content=content,
                filename=f"{safe_stem or 'embedded-image'}-{len(self.images) + 1}{extension}",
                mime_type=match.group(1),
                order=len(self.images),
                section_title=self._current_section,
            )
        )


def _mime_for_filename(filename: str) -> str | None:
    return {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".webp": "image/webp",
    }.get(Path(filename).suffix.casefold())
